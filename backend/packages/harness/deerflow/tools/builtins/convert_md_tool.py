"""Tool to convert Markdown files to PDF or DOCX.

Resolves sandbox virtual paths via Runtime context.
Handles inline Markdown: bold, italic, code, links, strikethrough.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Annotated

from langchain.tools import InjectedToolCallId, tool
from langchain_core.messages import ToolMessage
from langgraph.types import Command

from deerflow.tools.types import Runtime

logger = logging.getLogger(__name__)


@tool("convert_md", parse_docstring=True)
def convert_md_tool(
    runtime: Runtime,
    md_path: str,
    output_format: str = "pdf",
    output_path: str | None = None,
    tool_call_id: Annotated[str, InjectedToolCallId] = "",
) -> Command:
    """Convert a Markdown file to PDF or DOCX.

    The file must exist at the given sandbox path. Converts in-place.
    After conversion, use present_files to deliver the result to the user.

    Handles inline Markdown: **bold**, *italic*, `code`, [links](url), ~~strikethrough~~,
    code blocks (```), horizontal rules.

    Args:
        md_path: Absolute /mnt/user-data virtual path to the Markdown file.
        output_format: "pdf" or "docx". Defaults to "pdf".
        output_path: Optional output path. Defaults to same name with .pdf or .docx.

    Returns:
        Status message with the output path or error.
    """
    from deerflow.sandbox.tools import (
        get_thread_data,
        resolve_and_validate_user_data_path,
        validate_local_tool_path,
    )

    thread_data = get_thread_data(runtime)

    try:
        validate_local_tool_path(md_path, thread_data, read_only=True)
        real_md = Path(resolve_and_validate_user_data_path(md_path, thread_data))
    except Exception as e:
        return Command(
            update={"messages": [ToolMessage(f"Error: Cannot resolve path: {e}", tool_call_id=tool_call_id)]},
        )

    if not real_md.exists():
        return Command(
            update={"messages": [ToolMessage(f"Error: File not found at {real_md}", tool_call_id=tool_call_id)]},
        )

    if output_path is None:
        stem = Path(md_path).stem
        parent_virtual = Path(md_path).parent
        output_path = str(parent_virtual / f"{stem}.{output_format}")

    try:
        validate_local_tool_path(output_path, thread_data, read_only=False)
        real_output = Path(resolve_and_validate_user_data_path(output_path, thread_data))
    except Exception as e:
        return Command(
            update={"messages": [ToolMessage(f"Error: Cannot resolve output path: {e}", tool_call_id=tool_call_id)]},
        )

    real_output.parent.mkdir(parents=True, exist_ok=True)

    try:
        content = real_md.read_text(encoding="utf-8")
    except Exception as e:
        return Command(
            update={"messages": [ToolMessage(f"Error reading file: {e}", tool_call_id=tool_call_id)]},
        )

    if output_format == "pdf":
        result = _to_pdf(content, str(real_output))
    elif output_format == "docx":
        result = _to_docx(content, str(real_output))
    else:
        result = f"Unsupported format: {output_format}. Use 'pdf' or 'docx'."

    if result.startswith("OK:"):
        result = result.replace("OK:", f"Done. Saved to {output_path}")

    return Command(
        update={"messages": [ToolMessage(result, tool_call_id=tool_call_id)]},
    )


# --- Inline Markdown parsing ---

_INLINE_PATTERNS = [
    # Bold + italic ***text***
    (re.compile(r"\*\*\*(.+?)\*\*\*"), "bold_italic"),
    # Bold **text**
    (re.compile(r"\*\*(.+?)\*\*"), "bold"),
    # Italic *text* (single asterisk, not at start of line = list marker)
    (re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)"), "italic"),
    # Inline code `text`
    (re.compile(r"`([^`]+)`"), "code"),
    # Strikethrough ~~text~~
    (re.compile(r"~~(.+?)~~"), "strikethrough"),
    # Links [text](url)
    (re.compile(r"\[([^\]]+)\]\(([^)]+)\)"), "link"),
]


def _parse_inline(text: str) -> list:
    """Parse inline Markdown into a list of (text, formatting) tuples.

    Returns list of [{text, bold, italic, code, strikethrough, url}, ...]
    """
    results = [{"text": text}]
    for pattern, fmt in _INLINE_PATTERNS:
        new_results = []
        for part in results:
            if any(k in part for k in ("bold", "italic", "code", "link")):
                new_results.append(part)
                continue
            matches = list(pattern.finditer(part["text"]))
            if not matches:
                new_results.append(part)
                continue
            cursor = 0
            for m in matches:
                start, end = m.start(), m.end()
                if start > cursor:
                    new_results.append({"text": part["text"][cursor:start]})
                groups = m.groups()
                if fmt == "link":
                    new_results.append({"text": groups[0], "url": groups[1]})
                elif fmt == "bold":
                    new_results.append({"text": groups[0], "bold": True})
                elif fmt == "italic":
                    new_results.append({"text": groups[0], "italic": True})
                elif fmt == "bold_italic":
                    new_results.append({"text": groups[0], "bold": True, "italic": True})
                elif fmt == "code":
                    new_results.append({"text": groups[0], "code": True})
                elif fmt == "strikethrough":
                    new_results.append({"text": groups[0], "strikethrough": True})
                cursor = end
            if cursor < len(part["text"]):
                new_results.append({"text": part["text"][cursor:]})
        results = new_results
    return results


def _apply_formatting(doc, items: list) -> None:
    """Add a paragraph with inline formatting from parsed items."""
    p = doc.add_paragraph()
    for item in items:
        text = item.get("text", "")
        if not text:
            continue
        run = p.add_run(text)
        if item.get("bold"):
            run.bold = True
        if item.get("italic"):
            run.italic = True
        if item.get("code"):
            run.font.name = "Courier New"
            run.font.size = 95000  # 9.5pt
        if item.get("strikethrough"):
            run.font.strike = True
        if item.get("url"):
            run.font.color.rgb = None  # will set hyperlink style below
            try:
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                r_id = doc.part.relate_to(item["url"], RT.HYPERLINK, is_external=True)
                rPr = run._r.get_or_add_rPr()
                cS = parse_xml(f'<w:color {nsdecls("w")} w:val="0000FF"/>')
                rPr.append(cS)
                uS = parse_xml(f'<w:u {nsdecls("w")} w:val="single"/>')
                rPr.append(uS)
                # Link to actual URL via the run's r:id
                hyplink = parse_xml(
                    f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}>'
                    f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:hyperlink>'
                )
                # Replace the run's parent with the hyperlink
                p._p.append(hyplink)
                p._p.remove(run._r)
            except Exception:
                pass


# --- PDF converter (simple, no inline formatting) ---

def _to_pdf(content: str, output_path: str) -> str:
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        in_code_block = False
        for line in content.split("\n"):
            s = line.strip()
            if s.startswith("```"):
                in_code_block = not in_code_block
                if in_code_block:
                    pdf.set_font("Courier", size=9)
                    continue
                else:
                    pdf.set_font("Helvetica", size=11)
                    continue
            if in_code_block:
                pdf.set_font("Courier", size=9)
                pdf.multi_cell(0, 4, s)
                continue
            if not s:
                pdf.ln(4)
            elif s == "---" or s == "***":
                pdf.add_page()
            elif s and len(s) >= 3 and all(c == "_" for c in s):
                pdf.add_page()
            elif s == "<!-- pagebreak -->":
                pdf.add_page()
            elif s.startswith("# ") and not s.startswith("## "):
                pdf.set_font("Helvetica", "B", 16)
                pdf.cell(0, 10, _strip_inline(s[2:]), new_x="LMARGIN", new_y="NEXT")
            elif s.startswith("## ") and not s.startswith("### "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.cell(0, 8, _strip_inline(s[3:]), new_x="LMARGIN", new_y="NEXT")
            elif s.startswith("### "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 7, _strip_inline(s[4:]), new_x="LMARGIN", new_y="NEXT")
            elif s.startswith("- ") or s.startswith("* "):
                pdf.set_font("Helvetica", size=11)
                pdf.cell(5)
                pdf.multi_cell(0, 6, _strip_inline(s))
            else:
                pdf.set_font("Helvetica", size=11)
                pdf.multi_cell(0, 6, _strip_inline(s))

        pdf.output(output_path)
        return f"OK:{output_path}"
    except ImportError:
        return "fpdf2 is not installed. Run: cd backend && uv add fpdf2"
    except Exception as e:
        return f"PDF generation failed: {e}"


def _strip_inline(text: str) -> str:
    """Remove Markdown formatting markers for PDF (fpdf2 can't do inline)."""
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\s)\*(.+?)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


# --- DOCX converter (with inline formatting) ---

def _to_docx(content: str, output_path: str) -> str:
    try:
        from docx import Document
        from docx.shared import RGBColor, Pt

        doc = Document()

        in_code_block = False
        code_lines = []

        for line in content.split("\n"):
            s = line.strip()

            # Code blocks
            if s.startswith("```"):
                if in_code_block and code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9.5)
                    code_lines = []
                in_code_block = not in_code_block
                continue
            if in_code_block:
                code_lines.append(line.rstrip())
                continue

            # Empty lines
            if not s:
                continue

            # Page breaks
            if s == "---" or s == "***":
                doc.add_page_break()
                continue
            if len(s) >= 3 and all(c == "_" for c in s):
                doc.add_page_break()
                continue
            if s == "<!-- pagebreak -->":
                doc.add_page_break()
                continue

            # Headings
            if s.startswith("###### "):
                _apply_formatting(doc, _parse_inline(s[7:]), heading_level=6)
                continue
            if s.startswith("##### "):
                _apply_formatting(doc, _parse_inline(s[6:]), heading_level=5)
                continue
            if s.startswith("#### "):
                _apply_formatting(doc, _parse_inline(s[5:]), heading_level=4)
                continue
            if s.startswith("### ") and not s.startswith("#### "):
                _apply_formatting(doc, _parse_inline(s[4:]), heading_level=3)
                continue
            if s.startswith("## ") and not s.startswith("### "):
                _apply_formatting(doc, _parse_inline(s[3:]), heading_level=2)
                continue
            if s.startswith("# ") and not s.startswith("## "):
                _apply_formatting(doc, _parse_inline(s[2:]), heading_level=1)
                continue

            # Bullet lists
            if s.startswith("- ") or s.startswith("* "):
                bullet_text = s[2:] if s.startswith("- ") else s[1:].strip()
                _apply_formatting(doc, _parse_inline(bullet_text), bullet=True)
                continue

            # Regular paragraph
            _apply_formatting(doc, _parse_inline(s))

        doc.save(output_path)
        return f"OK:{output_path}"
    except ImportError:
        return "python-docx is not installed. Run: cd backend && uv add python-docx"
    except Exception as e:
        return f"DOCX generation failed: {e}"


def _apply_formatting(doc, items: list, heading_level: int | None = None, bullet: bool = False) -> None:
    """Add a paragraph with inline formatting from parsed items."""
    if heading_level:
        p = doc.add_heading("", level=heading_level)
    elif bullet:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()

    for item in items:
        text = item.get("text", "")
        if not text:
            continue
        run = p.add_run(text)

        if item.get("bold"):
            run.bold = True
        if item.get("italic"):
            run.italic = True
        if item.get("code"):
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        if item.get("strikethrough"):
            run.font.strike = True

        # Hyperlinks
        if item.get("url"):
            try:
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                r_id = doc.part.relate_to(item["url"], RT.HYPERLINK, is_external=True)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            except Exception:
                pass
